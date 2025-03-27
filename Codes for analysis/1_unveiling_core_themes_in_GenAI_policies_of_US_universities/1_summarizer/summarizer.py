from langchain_community.document_loaders import WebBaseLoader
from openai import OpenAI
import requests
from PyPDF2 import PdfReader
from langchain.document_loaders import WebBaseLoader
from io import BytesIO
import re
import os
from langchain.docstore.document import Document
import requests
import os



import os
import requests
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document  # Import the necessary library for .docx files
from langchain.document_loaders import WebBaseLoader

# Function to download PDF or DOCX with User-Agent header
def download_file_with_user_agent(url, local_filename):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    
    try:
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            with open(local_filename, 'wb') as f:
                f.write(response.content)
            print(f"File successfully downloaded: {local_filename}")
            return True
        else:
            print(f"Failed to retrieve the file. Status code: {response.status_code}")
            return False
    except Exception as e:
        print(f"Error downloading the file: {e}")
        return False

# Load documents from URLs including PDF, DOCX, and web pages
def load_documents(urls):
    texts = []
    
    for url in urls:
        if url.endswith('.pdf'):
            local_filename = 'temp_downloaded_file.pdf'  # Temporary file to store the downloaded PDF
            try:
                # Try to directly fetch the PDF content
                response = requests.get(url)
                if response.status_code == 200:
                    pdf_reader = PdfReader(BytesIO(response.content))
                    pdf_text = ""
                    for page in range(len(pdf_reader.pages)):
                        pdf_text += pdf_reader.pages[page].extract_text() + "\n"
                    texts.append(pdf_text)
                else:
                    print(f"Failed to retrieve the PDF directly from {url}, trying with user-agent download.")
                    # If direct download fails, try using the user-agent download function
                    if download_file_with_user_agent(url, local_filename):
                        # Read the downloaded PDF
                        with open(local_filename, 'rb') as f:
                            pdf_reader = PdfReader(f)
                            pdf_text = ""
                            for page in range(len(pdf_reader.pages)):
                                pdf_text += pdf_reader.pages[page].extract_text() + "\n"
                            texts.append(pdf_text)
                        # After processing, delete the temporary file
                        os.remove(local_filename)
                    else:
                        print(f"Failed to download the PDF from {url} with user-agent.")
            except Exception as e:
                print(f"Error loading PDF from {url}: {e}")
        
        elif url.endswith('.docx'):
            local_filename = 'temp_downloaded_file.docx'  # Temporary file to store the downloaded DOCX
            try:
                # Try to directly fetch the DOCX content
                response = requests.get(url)
                if response.status_code == 200:
                    docx_content = BytesIO(response.content)
                    document = Document(docx_content)
                    docx_text = "\n".join([para.text for para in document.paragraphs])
                    texts.append(docx_text)
                else:
                    print(f"Failed to retrieve the DOCX directly from {url}, trying with user-agent download.")
                    # If direct download fails, try using the user-agent download function
                    if download_file_with_user_agent(url, local_filename):
                        # Read the downloaded DOCX
                        document = Document(local_filename)
                        docx_text = "\n".join([para.text for para in document.paragraphs])
                        texts.append(docx_text)
                        # After processing, delete the temporary file
                        os.remove(local_filename)
                    else:
                        print(f"Failed to download the DOCX from {url} with user-agent.")
            except Exception as e:
                print(f"Error loading DOCX from {url}: {e}")
        
        else:
            try:
                # Use WebBaseLoader for non-PDF/DOCX documents
                loader = WebBaseLoader(url)
                web_texts = loader.load()
                texts.append(" ".join([item.page_content for item in web_texts]))
            except Exception as e:
                print(f"Error loading non-PDF/DOCX content from {url}: {e}")

    return texts



def clean_text_list(text_list):
    cleaned_list = []
    
    for text in text_list:
        # Step 1: Replace multiple newlines with just two '\n'
        cleaned_text = re.sub(r'\n\s*\n+', ' \n\n ', text)
        
        # Step 2: Replace '\xa0' with a regular space
        cleaned_text = cleaned_text.replace('\xa0', ' ')
        
        # Append the cleaned text to the list
        cleaned_list.append(cleaned_text)
    
    return cleaned_list

def get_message(text):
    system_message = '''You are a formal and professional assistant helping a higher education administrator summarize complex AI policies.
    Your responses should be concise, accurate, and formatted in clear, bulleted key points.'''

    # Task prompt for the model
    task_prompt = '''Your role is to synthesize and summarize the university's AI policies based on the provided document into 500 words or less.
    Focus on identifying key topics and summarizing them in a succint, clear and concise manner. 
    Present the information in a numbered list of key points, ensuring each point is relevant to the document's discussion of AI policy.
    Ensure your summary succintly captures the most important themes, keeping each key point brief and informative.
    Avoid including any minor points or outliers. '''


    Example_1='''Dear Members of the WashU Community, 
    There has been much discussion recently surrounding the use of Generative Artificial Intelligence (AI), AI that uses algorithms to quickly produce content such as text, images, music, videos, code, or other media through the use of commands or prompts.  
    The power and functionality of these tools has great potential to be incredibly beneficial across all disciplines and among all organizations. Examples of such AI tools include Machine Learning (ML) and Large Language Models (LLMs.) 
    Because this field is evolving so quickly, it is important for our IT team to be at the forefront of its application so that everyone can work securely, effectively, and efficiently.  
    To equip you for success, we would like to offer a brief explanation of each of these tools, as well as introduce guidelines on the use of generative AI, such as OpenAI’s ChatGPT, Google Bard, and many others. 
    AI is a machine’s ability to perform a task that would normally require human intelligence. ML and LLMs leverage AI to give machines the ability to adapt, or to compile massive amounts of information used to replicate human writing, speech, and behavior.  
    The university supports and encourages the responsible and secure exploration of AI tools. When using any of these tools, especially public, open-source, non-protected AI tools, it is vitally important to keep information security and data privacy, compliance, copyright, and academic integrity in mind. Currently, we are exploring privacy compliant LLM solutions. Those with a particular need in this area, or questions, are encouraged to reach out for a consultation via aiquestions@wustl.edu.  
    It is clear AI is a rapidly evolving technology. The university is tracking developments and adapting plans to support the community in secure, compliant, and privacy-respecting ways. Guidelines will undergo updates to match advancements and innovation with proper safety controls.
    Initial guidelines for use of AI tools: 
    Be mindful not to share sensitive information: Please do not enter confidential or protected data or information, including non-public research data, into publicly available or vendor-enabled AI tools. Information shared with public AI tools is usually claimed to be the property of the vendor, is not considered private, and could expose proprietary or sensitive information to unauthorized parties. It is the user’s responsibility to protect confidential data.
    These tools can be inaccurate: Each individual is responsible for any content that is produced or published containing AI-generated material.Note that AI tools sometimes “hallucinate,” generating content that can be highly convincing, but inaccurate, misleading, or entirely fabricated. Furthermore, it may contain copyrighted material. It is imperative that all AI-generated content be reviewed carefully for correctness before submission or publication. It is the user’s responsibility to verify everything.
    Adhere to current academic integrity policies: Review university, school and department handbooks and policies for student and faculty. Schools will be developing and updating their policies as we learn more about AI tools. Faculty members should make clear to students they are teaching and advising about policies on the permitted uses, if any, of AI in classes and on academic work. Students are also encouraged to ask their instructors for clarification about these policies as needed.
    Be alert for AI-enabled phishing: AI has made it easier for malicious actors to create sophisticated scams at a far greater scale. Continue to follow security best practices and report suspicious messages via the Phish Report button in Outlook or to phishing@wustl.edu.
    Connect with WashU IT before procuring generative AI tools: The university is working to ensure that tools procured on behalf of WashU have the appropriate privacy and security protections. 
    If you have procured or are considering procuring AI tools, or if you have questions, please contact WashU IT at aiquestions@wustl.edu.  
    Vendor generative AI tools must be assessed for risk by WashU’s Office of Information Security prior to use. This includes the following: 
    Existing tools that add or expand AI capabilities. 
    New purchases of vendor AI tools. 
    It is important to note that these guidelines are not new university policy; rather, they leverage existing university policies. Please look for institution-specific guidance to follow this communication. We look forward to working with you in the spirit of collaboration and innovation. 
    
    Sincerely, 
    Chris Shull 
    Chief Information Security Officer 
    Gregory Hart 
    Chief Technology Officer'''

    output_1='''1. Security, Privacy, and Compliance in the use of AI tools: 
    WashU places a heavy emphasis on using AI technologies in a manner that is secure, compliant, and privacy-respecting. These include sharing sensitive information, 
    being alert for AI-enabled phishing and consulting WashU IT before procuring generative AI tools, which include 
    Vendor generative AI tools must be assessed for risk by WashU’s Office of Information Security prior to using existing AI tools. 

    
    2. Responsible use of AI:
    WashU encourages careful review of AI-generated content due to the potential for inaccuracies or "hallucinations"
    by AI models. It reinforces the user’s responsibility to verify AI-generated material before publication or submission.
    

    3. Adherence to academic integrity policies:
    WashU adherence to academic integrity policies, with an emphasis on developing clear guidelines around the use of AI in academic work, 
    and ensuring students and faculty are informed of these policies. Individuals are encouraged to review university, 
    school and department handbooks and policies for student and faculty. 

    '''

    Example_2='''
    Per the Graduate Bulletin, the master’s thesis demonstrates independent judgment in developing a problem from primary sources, and a dissertation represents originality in research, independent thinking, scholarly ability, and technical mastery of a field of study. It is the responsibility of the advisory committee to review and evaluate the thesis or dissertation as a representation of a student’s individual effort. As such, the use of generative AI in theses and dissertations is considered unauthorized assistance per the Academic Code of Honesty and is prohibited unless specifically authorized by members of the advisory committee for use within the approved scope. If approved by the advisory committee, the extent of generative AI usage should be disclosed in a statement within the thesis or dissertation.

    Guidance from Academic Honesty: honesty.uga.edu/Academic-Honesty-Policy/Prohibited_Conduct/

    Giving or receiving help for assignments without prior approval from your instructor. During any assignment, any help (such as books, notes, calculators, technology, internet resources, or conversations with others) is considered unauthorized unless the instructor explicitly allows it. Examples include, but are not limited to:

    Copying, or allowing others to copy, answers to an assignment.
    Sending, receiving, posting, uploading, downloading, or accessing relevant exam information, prior to, during, or after the exam itself (including written or orally, or use of sign, electronic device, or digital resource information).
    Completing someone else’s assignment or allowing them to complete yours.
    Collaborating on any assignment that is an individual assignment.
    Submitting group work that does not represent work from all members of the group. Every student whose name is on a group project is responsible for the academic honesty of the group assignment.
    Using any cellular device, electronic device, digital device, or programmable calculator without permission during an exam or closed assignment.
    The bottom line:

    If you are requesting, sharing, or receiving any assignment or test information and it is an individual assignment, you are putting yourself at risk.
    The whole group is responsible for the integrity of group work.
    Don’t access any electronic devices or notes for any reason unless your instructor explicitly says it’s allowed during an exam.
    Never use Artificial Intelligence on an assignment unless it is explicitly authorized by your instructor before the assignment is turned in.
    
    '''

    Output_2='''1. Prohibition of Unauthorized Use of Generative AI in Theses and Dissertations:
    Generative AI is considered unauthorized assistance unless specifically authorized by the student's advisory committee. If approved for use, the extent of AI usage must be disclosed in a statement within the thesis or dissertation.

    
    2. General Academic Honesty Regarding AI Usage:
    Use of any technology, including generative AI, for assignments is prohibited unless explicitly allowed by the instructor.
    Unauthorized use of AI to assist in completing assignments, exams, or tests is considered a violation of academic honesty policies.
    Students must not use AI on assignments unless authorized before submission, and all forms of AI-assisted work must follow the instructor's guidelines.'''

    prompt=f"Summarize the following AI policy, adhering to the few-shot examples provided: {text}"
    message=[
            {
                "role": "system",
                "content": system_message
            },
            {
                "role": "user",
                "content": task_prompt
            },
            {
                "role": "user",  # Example user input
                "content": Example_1  # Optional, can remove this part to avoid bias
            },
            {
                "role": "assistant",  # Example assistant output
                "content": output_1  # Optional, can remove this part to avoid bias
            },
            {
                "role": "user",  # Example user input
                "content": Example_2  # Optional, can remove this part to avoid bias
            },
            {
                "role": "assistant",  # Example assistant output
                "content": Output_2  # Optional, can remove this part to avoid bias
            },
            {
                "role": "user",  # Now pass in the actual prompt
                "content": prompt
            }
        ]
    return(message)

def generate_response_openAI(text,api_key, model = "gpt-4o-mini", max_tokens = 500):
    """
    Function to get a response from the OpenAI GPT-4 model.

    Parameters:
    prompt (str): The prompt that will be sent to the model for completion.
    api_key (str): The API key for the OpenAI API.
    model (str, optional): The model to be used for completion.
    Default to "gpt-4".
    max_tokens (int, optional): The maximum length of the generated text.
    Defaults to 1000.

    Returns:
    str: The generated text from the model.
    """
    messages=get_message(text)
    client = OpenAI(
    api_key=api_key,
    )


    response = client.chat.completions.create(
        model = model,
        messages=messages,
        max_tokens=max_tokens, 
        temperature=0.1
    )
    return response.choices[0].message.content

