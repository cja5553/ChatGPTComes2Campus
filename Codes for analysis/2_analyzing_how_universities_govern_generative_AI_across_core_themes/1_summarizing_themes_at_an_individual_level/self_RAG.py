# self_RAG.py

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader
from langchain_openai import OpenAIEmbeddings
from langchain.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain.retrievers.multi_query import MultiQueryRetriever
from langgraph.graph import StateGraph, START, END
from pydantic import BaseModel, Field
from pprint import pprint
from langchain.retrievers import ParentDocumentRetriever, ContextualCompressionRetriever, MultiQueryRetriever
from langchain.storage import InMemoryStore
from langchain import hub
from typing import List
from typing_extensions import TypedDict
from langchain_openai import ChatOpenAI
from pprint import pprint

from langchain.schema import Document  # Assuming you're using Langchain's Document schema
from PyPDF2 import PdfReader
import requests
from io import BytesIO
import os
from docx import Document as DocxDocument
from langchain_chroma import Chroma
import re
from chromadb import PersistentClient
from uuid import uuid4


api_key = "<insert openAI API key>"
os.environ['OPENAI_API_KEY'] = api_key

# Function to clean the extracted text content
def clean_text_list(text_list):
    cleaned_list = []
    
    for text in text_list:
        # Step 1: Replace multiple newlines with just two '\n'
        cleaned_text = re.sub(r'\n\s*\n+', ' \n\n ', text)
        
        # Step 2: Replace '\xa0' with a regular space
        cleaned_text = cleaned_text.replace('\xa0', ' ')
        
        # Append the cleaned text to the list
        cleaned_list.append(cleaned_text)
    
    return cleaned_list

# Function to download PDF or DOCX with User-Agent header
def download_file_with_user_agent(url, local_filename):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    
    try:
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            with open(local_filename, 'wb') as f:
                f.write(response.content)
            print(f"File successfully downloaded: {local_filename}")
            return True
        else:
            print(f"Failed to retrieve the file. Status code: {response.status_code}")
            return False
    except Exception as e:
        print(f"Error downloading the file: {e}")
        return False

# Load documents from URLs including PDF, DOCX, and web pages
def load_documents(urls):
    docs = []
    
    for url in urls:
        if url.endswith('.pdf'):
            local_filename = 'temp_downloaded_file.pdf'  # Temporary file to store the downloaded PDF
            try:
                # Try to directly fetch the PDF content
                response = requests.get(url)
                if response.status_code == 200:
                    pdf_reader = PdfReader(BytesIO(response.content))
                    pdf_text = ""
                    for page in range(len(pdf_reader.pages)):
                        pdf_text += pdf_reader.pages[page].extract_text() + "\n"
                    
                    # Clean the extracted text
                    cleaned_text = clean_text_list([pdf_text])[0]
                    
                    # Append cleaned text as a Document
                    docs.append(Document(page_content=cleaned_text, metadata={"source": url}))
                else:
                    print(f"Failed to retrieve the PDF directly from {url}, trying with user-agent download.")
                    if download_file_with_user_agent(url, local_filename):
                        with open(local_filename, 'rb') as f:
                            pdf_reader = PdfReader(f)
                            pdf_text = ""
                            for page in range(len(pdf_reader.pages)):
                                pdf_text += pdf_reader.pages[page].extract_text() + "\n"
                            
                            # Clean the extracted text
                            cleaned_text = clean_text_list([pdf_text])[0]
                            
                            docs.append(Document(page_content=cleaned_text, metadata={"source": url}))
                        os.remove(local_filename)
                    else:
                        print(f"Failed to download the PDF from {url} with user-agent.")
            except Exception as e:
                print(f"Error loading PDF from {url}: {e}")
        
        elif url.endswith('.docx'):
            local_filename = 'temp_downloaded_file.docx'  # Temporary file to store the downloaded DOCX
            try:
                # Try to directly fetch the DOCX content
                response = requests.get(url)
                if response.status_code == 200:
                    docx_content = BytesIO(response.content)
                    document = DocxDocument(docx_content)
                    docx_text = "\n".join([para.text for para in document.paragraphs])
                    
                    # Clean the extracted text
                    cleaned_text = clean_text_list([docx_text])[0]
                    
                    docs.append(Document(page_content=cleaned_text, metadata={"source": url}))
                else:
                    print(f"Failed to retrieve the DOCX directly from {url}, trying with user-agent download.")
                    if download_file_with_user_agent(url, local_filename):
                        document = DocxDocument(local_filename)
                        docx_text = "\n".join([para.text for para in document.paragraphs])
                        
                        # Clean the extracted text
                        cleaned_text = clean_text_list([docx_text])[0]
                        
                        docs.append(Document(page_content=cleaned_text, metadata={"source": url}))
                        os.remove(local_filename)
                    else:
                        print(f"Failed to download the DOCX from {url} with user-agent.")
            except Exception as e:
                print(f"Error loading DOCX from {url}: {e}")
        
        else:
            try:
                # Use WebBaseLoader for non-PDF/DOCX documents
                loader = WebBaseLoader(url)
                web_texts = loader.load()
                web_contents = [item.page_content for item in web_texts]
                
                # Clean the extracted text
                cleaned_texts = clean_text_list(web_contents)
                
                # Append cleaned text to docs
                for i, item in enumerate(web_texts):
                    docs.append(Document(page_content=cleaned_texts[i], metadata={"source": url}))
            except Exception as e:
                print(f"Error loading non-PDF/DOCX content from {url}: {e}")

    return docs




# Text splitter for document chunking
def split_documents(docs_list, chunk_size=250, chunk_overlap=50):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return text_splitter.split_documents(docs_list)






# Grading Models
class GradeDocuments(BaseModel):
    """Binary score for relevance check on retrieved documents."""

    binary_score: str = Field(
        description="Documents are relevant to the question, 'yes' or 'no'"
    )


class GradeHallucinations(BaseModel):
    """Binary score for hallucination present in generation answer."""

    binary_score: str = Field(
        description="Answer is grounded in the facts, 'yes' or 'no'"
    )

class GradeAnswer(BaseModel):
    """Binary score to assess answer addresses question."""

    binary_score: str = Field(
        description="Answer addresses the question, 'yes' or 'no'"
    )



class GraphState(TypedDict):
    """
    Represents the state of our graph.

    Attributes:
        question: question
        generation: LLM generation
        documents: list of documents
    """

    question: str
    generation: str
    documents: List[str]


# Define retrieve function using Parent_Document_Retriever
def retrieve(state, retriever):
    print("---RETRIEVE---")
    question = state["question"]
    documents = retriever.get_relevant_documents(question)
    return {"documents": documents, "question": question}

def generate(state):
    """
    Generate answer

    Args:
        state (dict): The current graph state

    Returns:
        state (dict): New key added to state, generation, that contains LLM generation
    """
    print("---GENERATE---")
    question = state["question"]
    documents = state["documents"]

    # RAG generation
    prompt = hub.pull("rlm/rag-prompt")
    llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0, cache=False)
    rag_chain = prompt | llm | StrOutputParser()

    generation = rag_chain.invoke({"context": documents, "question": question})
    return {"documents": documents, "question": question, "generation": generation}



def grade_documents(state):
    """
    Determines whether the retrieved documents are relevant to the question.

    Args:
        state (dict): The current graph state

    Returns:
        state (dict): Updates documents key with only filtered relevant documents
    """

    print("---CHECK DOCUMENT RELEVANCE TO QUESTION---")
    question = state["question"]
    documents = state["documents"]

    
    # LLM with function call
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0, cache=False)
    structured_llm_grader = llm.with_structured_output(GradeDocuments)

    # Prompt
    system = """You are a grader assessing relevance of a retrieved document to a user question. \n 
        It does not need to be a stringent test. The goal is to filter out erroneous retrievals. \n
        If the document contains keyword(s) or semantic meaning related to the user question, grade it as relevant. \n
        Give a binary score 'yes' or 'no' score to indicate whether the document is relevant to the question."""
    grade_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", system),
            ("human", "Retrieved document: \n\n {document} \n\n User question: {question}"),
        ]
    )

    retrieval_grader = grade_prompt | structured_llm_grader
    # Score each doc
    filtered_docs = []
    for d in documents:
        score = retrieval_grader.invoke(
            {"question": question, "document": d.page_content}
        )
        grade = score.binary_score
        if grade == "yes":
            print("---GRADE: DOCUMENT RELEVANT---")
            filtered_docs.append(d)
        else:
            print("---GRADE: DOCUMENT NOT RELEVANT---")
            continue
    return {"documents": filtered_docs, "question": question}


# Define transform_query function
def transform_query(state):
    """
    Transform the query to produce a better question.

    Args:
        state (dict): The current graph state

    Returns:
        state (dict): Updates question key with a re-phrased question
    """

    print("---TRANSFORM QUERY---")
    question = state["question"]
    documents = state["documents"]
    ### Question Re-writer

    # LLM
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0, cache=False)

    # Prompt
    system = """You a question re-writer that converts an input question to a better version that is optimized \n 
        for vectorstore retrieval. Look at the input and try to reason about the underlying semantic intent / meaning."""
    re_write_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", system),
            (
                "human",
                "Here is the initial question: \n\n {question} \n Formulate an improved question.",
            ),
        ]
    )

    question_rewriter = re_write_prompt | llm | StrOutputParser()

    # Re-write question
    better_question = question_rewriter.invoke({"question": question})
    return {"documents": documents, "question": better_question}


def decide_to_generate(state):
    """
    Determines whether to generate an answer, or re-generate a question.

    Args:
        state (dict): The current graph state

    Returns:
        str: Binary decision for next node to call
    """

    print("---ASSESS GRADED DOCUMENTS---")
    state["question"]
    filtered_documents = state["documents"]

    if not filtered_documents:
        # All documents have been filtered check_relevance
        # We will re-generate a new query
        print(
            "---DECISION: ALL DOCUMENTS ARE NOT RELEVANT TO QUESTION, TRANSFORM QUERY---"
        )
        return "transform_query"
    else:
        # We have relevant documents, so generate answer
        print("---DECISION: GENERATE---")
        return "generate"
    


def grade_generation_v_documents_and_question(state):
    """
    Determines whether the generation is grounded in the document and answers question.

    Args:
        state (dict): The current graph state

    Returns:
        str: Decision for next node to call
    """

    print("---CHECK HALLUCINATIONS---")
    question = state["question"]
    documents = state["documents"]
    generation = state["generation"]


    # hallucination grader
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0, cache=False)
    structured_llm_grader = llm.with_structured_output(GradeHallucinations)

    # Prompt
    system = """
    You are a grader assessing whether an LLM generation is grounded in / supported by a set of retrieved facts. 

    Strictly give a binary score of either 'YES' or 'NO'. 
    - 'YES' means that the answer is mostly grounded in / supported by the set of facts, even if there are minor variations or minor unsupported details that do not affect the overall accuracy or intent of the answer.
    - 'NO' means you are certain that the generation is massively hallucinated or significantly incorrect. This means the generation contains critical, unsupported content or is highly misleading.

    Examples:
    - Set of facts: "The capital of France is Paris."
      LLM generation: "Paris is the capital of France, known for the Eiffel Tower."
      -> Score: 'YES' (Minor addition that doesn't contradict the facts)

    - Set of facts: "The capital of France is Paris."
      LLM generation: "Rome is the capital of France."
      -> Score: 'NO'
    """

    hallucination_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", system),
            ("human", "Set of facts: \n\n {documents} \n\n LLM generation: {generation}"),
        ]
    )

    hallucination_grader = hallucination_prompt | structured_llm_grader
    score = hallucination_grader.invoke(
        {"documents": documents, "generation": generation}
    )
    grade = score.binary_score

    # answer grader
    # LLM with function call
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0, cache=False)
    structured_llm_grader = llm.with_structured_output(GradeAnswer)

    # Prompt
    system = """You are a grader assessing whether an answer addresses / resolves a question \n 
        Give a binary score 'yes' or 'no'. Yes' means that the answer resolves the question."""
    answer_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", system),
            ("human", "User question: \n\n {question} \n\n LLM generation: {generation}"),
        ]
    )

    answer_grader = answer_prompt | structured_llm_grader
    print("Grade:", grade)
    # Check hallucination
    if grade == "yes" or grade == "YES" or grade == "Yes":
        print("---DECISION: GENERATION IS GROUNDED IN DOCUMENTS---")
        # Check question-answering
        print("---GRADE GENERATION vs QUESTION---")
        score = answer_grader.invoke({"question": question, "generation": generation})
        grade = score.binary_score
        if grade == "yes":
            print("---DECISION: GENERATION ADDRESSES QUESTION---")
            return "useful"
        else:
            print("---DECISION: GENERATION DOES NOT ADDRESS QUESTION---")
            return "not useful"
    else:
        pprint("---DECISION: GENERATION IS NOT GROUNDED IN DOCUMENTS, RE-TRY---")
        return "not supported"

# Helper function to delete all Chroma collections
def delete_all_chroma_collections():
    try:
        # Initialize the persistent Chroma client
        chroma_client = PersistentClient(path="./chroma_db/txt_db")
        
        # Get the list of all collections
        collections = chroma_client.list_collections()
        
        # Delete each collection
        for collection in collections:
            chroma_client.delete_collection(collection.name)
            print(f"Deleted collection: {collection.name}")
        
        if not collections:
            print("No collections to delete.")
    except Exception as e:
        raise Exception(f"Unable to delete collections: {e}")




# Function to build the workflow
def build_workflow(urls, retriever=None):
    # Delete all existing collections to start fresh
    delete_all_chroma_collections()

    # Initialize Chroma and splitters
    parent_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)  # For parent documents
    child_splitter = RecursiveCharacterTextSplitter(chunk_size=200, chunk_overlap=50)  # For child documents
    documents = load_documents(urls)
    print(len(urls), "DOCUMENTS LOADED, BEGINNING ANALYSIS...")

    # Check retriever type and build vectorstore accordingly
    if retriever == "MultiQueryRetriever":
        splits = parent_splitter.split_documents(documents)
        vectordb = Chroma.from_documents(
            documents=splits, 
            embedding=OpenAIEmbeddings(model="text-embedding-3-small"),
            persist_directory="./chroma_db/txt_db"  # Use the specified directory
        )
        retriever = MultiQueryRetriever.from_llm(
            retriever=vectordb.as_retriever(), 
            llm=ChatOpenAI(temperature=0)
        )
    elif retriever == "ParentDocumentRetriever":
        vectorstore = Chroma.from_documents(
            documents=documents,
            persist_directory="./chroma_db/txt_db",  # Use the specified directory
            embedding=OpenAIEmbeddings(model="text-embedding-3-small")
        )
        retriever = ParentDocumentRetriever(
            vectorstore=vectorstore,
            docstore=InMemoryStore(),  # Initialize an in-memory store
            child_splitter=child_splitter,
            parent_splitter=parent_splitter
        )
        retriever.add_documents(documents)  # Add documents to the retriever
    else:
        doc_splits = parent_splitter.split_documents(documents)
        # Add to vectorDB
        vectorstore = Chroma.from_documents(
            documents=doc_splits,
            persist_directory="./chroma_db/txt_db",  # Use the specified directory
            embedding=OpenAIEmbeddings(model="text-embedding-3-small"),
        )
        retriever = vectorstore.as_retriever()

    # Log the count of documents in the collection
    print("Document count in vectorstore after adding new documents:", vectorstore._collection.count())

    # Build the workflow using the retrieved documents
    workflow = StateGraph(GraphState)

    # Define the nodes in the workflow
    workflow.add_node("retrieve", lambda state: retrieve(state, retriever))  # retrieve
    workflow.add_node("grade_documents", lambda state: grade_documents(state))  # grade documents
    workflow.add_node("generate", lambda state: generate(state))  # generate
    workflow.add_node("transform_query", lambda state: transform_query(state))  # transform_query

    # Build graph
    workflow.add_edge(START, "retrieve")
    workflow.add_edge("retrieve", "grade_documents")
    workflow.add_conditional_edges(
        "grade_documents",
        decide_to_generate,
        {
            "transform_query": "transform_query",
            "generate": "generate",
        },
    )
    workflow.add_edge("transform_query", "retrieve")
    workflow.add_conditional_edges(
        "generate",
        grade_generation_v_documents_and_question,
        {
            "not supported": "generate",
            "useful": END,
            "not useful": "transform_query",
        },
    )

    # Compile the workflow
    return workflow.compile()

# Function to execute the self-RAG system
def execute_self_RAG(urls, question, retriever=None, print_statement=False):
    # Build the workflow
    app = build_workflow(urls, retriever=retriever)
    
    # Set the inputs for the workflow
    inputs = {"question": question}
    
    # Run the workflow and capture the output
    for output in app.stream(inputs):
        for key, value in output.items():
            # Optional: Print state at each node
            pprint(f"Node '{key}':")
            if print_statement:
                pprint(value)
        pprint("\n---\n")

    # Return the final generated answer
    return value["generation"]



